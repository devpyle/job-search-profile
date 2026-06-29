#!/usr/bin/env python3
"""
Job Search Dashboard — Kanban board for tracking applications + resume/CL generation.

Run:  python3 scripts/dashboard.py
Open: http://localhost:5000

Dependencies:
    pip install flask python-docx anthropic requests python-dotenv
"""

import hashlib
import html
import json
import os
import re
import sqlite3
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

import anthropic
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv
from flask import (Flask, g, jsonify, redirect, render_template,
                   request, send_file, url_for)

load_dotenv()

from startup import validate
validate(
    env_optional={"ANTHROPIC_API_KEY": "AI-powered document generation"},
    config_attrs=["CANDIDATE_NAME", "JOB_DOCS", "HOME_METRO_TERMS", "HOME_CITY"],
    script_name="dashboard.py",
)

# ── PATHS ─────────────────────────────────────────────────────────────────────

REPO_ROOT   = Path(__file__).parent.parent
DOCS_DIR    = REPO_ROOT / "docs"
OUTPUT_DIR  = REPO_ROOT / "output" / "job-radar"
DASH_DIR    = REPO_ROOT / "dashboard"
DB_PATH     = DASH_DIR / "data" / "jobs.db"
EXPORT_DIR  = REPO_ROOT / "output" / "documents"
INTERVIEW_PREP_DIR = REPO_ROOT / "output" / "interview-prep"

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# Env for the `claude` CLI: strip ANTHROPIC_API_KEY so the CLI uses the
# Claude.ai (OAuth / Max subscription) login instead of billing API credits.
# Leaving the key set both bills per-token credits and triggers the
# "connectors are disabled" error that breaks `claude -p`.
_CLI_ENV = {k: v for k, v in os.environ.items() if k != "ANTHROPIC_API_KEY"}

# ── PERSONAL CONFIG (from config.py) ──────────────────────────────────────────
sys.path.insert(0, str(REPO_ROOT))
from config import CANDIDATE_NAME, JOB_DOCS, HOME_METRO_TERMS, HOME_CITY  # noqa: E402
from portal_scanner import scan_all as portal_scan_all  # noqa: E402
import db as data  # noqa: E402

# ── KANBAN COLUMNS ────────────────────────────────────────────────────────────

STATUSES = [
    "New", "Reviewing", "Drafting", "Ready",
    "Applied", "Interviewing", "Offer",
    "Accepted", "Rejected", "Passed",
]
APP_STATUSES       = ["New", "Reviewing", "Drafting", "Ready", "Applied"]
INTERVIEW_STATUSES = ["Interviewing", "Offer"]
END_STATUSES       = ["Accepted", "Rejected", "Passed"]
ACTIVE_STATUSES    = APP_STATUSES  # kept for any legacy references

STATUS_COLORS = {
    "New":          "#94a3b8",
    "Reviewing":    "#3b82f6",
    "Drafting":     "#8b5cf6",
    "Ready":        "#10b981",
    "Applied":      "#06b6d4",
    "Interviewing": "#f97316",
    "Offer":        "#84cc16",
    "Accepted":     "#22c55e",
    "Rejected":     "#ef4444",
    "Passed":       "#6b7280",
}

# ── FLASK APP ─────────────────────────────────────────────────────────────────

app = Flask(
    __name__,
    template_folder=str(DASH_DIR / "templates"),
    static_folder=str(DASH_DIR / "static"),
)
app.secret_key = "job-search-dashboard-local"


@app.after_request
def no_cache_html(response):
    # Pages are always rendered fresh from the DB / radar reports. Without this,
    # browsers heuristically cache the HTML and a normal refresh serves a stale
    # page (e.g. radar jobs appearing frozen on the day first loaded). Static
    # assets keep their default caching.
    if response.mimetype == "text/html":
        response.headers["Cache-Control"] = "no-store, must-revalidate"
    return response


# ── DATABASE ──────────────────────────────────────────────────────────────────

def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(str(DB_PATH))
        g.db.row_factory = sqlite3.Row
    return g.db

@app.teardown_appcontext
def close_db(error):
    db = g.pop("db", None)
    if db:
        db.close()

def init_db():
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(DB_PATH))
    data.init_schema(conn)
    conn.close()

# ── REPORT PARSER ─────────────────────────────────────────────────────────────

def job_id_from_url(url: str) -> str:
    return hashlib.sha1(url.encode()).hexdigest()[:8]

def list_reports() -> list[dict]:
    files = sorted(OUTPUT_DIR.glob("[0-9][0-9][0-9][0-9]-*.md"), reverse=True)
    reports = []
    for f in files:
        if f.name.startswith("."):
            continue
        parts = f.stem.split("-")
        if len(parts) >= 4:
            date_str = "-".join(parts[:3])
            slot     = parts[3].upper()
            reports.append({
                "filename": f.name,
                "date":     date_str,
                "slot":     slot,
                "path":     str(f),
            })
    return reports

def parse_report(filepath: str) -> list[dict]:
    """Parse a radar .md file into a list of job dicts."""
    text = Path(filepath).read_text(encoding="utf-8")

    fname       = Path(filepath).stem
    parts       = fname.split("-")
    report_date = "-".join(parts[:3]) if len(parts) >= 3 else ""
    report_slot = parts[3].upper()    if len(parts) >= 4 else ""
    report_file = Path(filepath).name

    TIER_LABELS = {
        # Tier label mapping
        "APPLY NOW":    "Apply Now",
        "WORTH A LOOK": "Worth a Look",
        "WEAK MATCH":   "Weak Match",
        "FILTERED OUT": "Skip",
        # Casey's tier names
        "PERFECT FIT":  "Apply Now",
        "GOOD FIT":     "Worth a Look",
        "SKIP":         "Skip",
    }

    jobs: list[dict] = []
    current_tier = "Unknown"
    in_filtered  = False

    # Split on the ─── divider lines
    sections = re.split(r"\n[─\-]{20,}\n", text)

    for section in sections:
        section = section.strip()
        if not section:
            continue

        # Detect tier header (short section with a known label)
        upper = section.upper()
        matched = None
        for key, val in TIER_LABELS.items():
            if key in upper and len(section) < 100:
                matched = val
                break
        if matched:
            current_tier = matched
            in_filtered  = matched in ("Skip", "Weak Match")
            continue

        if in_filtered:
            # Compact format:  ✗/⚠️ TITLE — COMPANY  [reason]\n   URL
            pending_job: Optional[dict] = None
            for line in section.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("✗") or line.startswith("⚠") or line.startswith("⛔"):
                    line = re.sub(r"^[✗⚠⛔️\s]+", "", line).strip()
                    reason_match = re.search(r"\[([^\]]+)\]\s*$", line)
                    reason = reason_match.group(1) if reason_match else ""
                    if reason_match:
                        line = line[:reason_match.start()].strip()
                    if " — " in line:
                        title_part, company_part = line.split(" — ", 1)
                    else:
                        title_part, company_part = line, ""
                    pending_job = {
                        "title": title_part.strip(), "company": company_part.strip(),
                        "location": "", "salary": "", "reason": reason,
                        "description": "", "posted": "", "source": "",
                        "url": "", "tier": current_tier,
                        "report_date": report_date, "report_slot": report_slot,
                        "report_file": report_file, "job_id": "",
                    }
                    jobs.append(pending_job)
                elif line.startswith("http") and jobs and not jobs[-1]["url"]:
                    jobs[-1]["url"]    = line
                    jobs[-1]["job_id"] = job_id_from_url(line)
        else:
            # Full format — jobs separated by blank lines
            blocks = re.split(r"\n\n+", section)
            for block in blocks:
                block = block.strip()
                if not block:
                    continue
                lines = block.split("\n")
                if len(lines) < 2:
                    continue

                job: dict = {
                    "title": "", "company": "", "location": "", "salary": "",
                    "reason": "", "description": "", "posted": "", "source": "",
                    "url": "", "tier": current_tier,
                    "report_date": report_date, "report_slot": report_slot,
                    "report_file": report_file, "job_id": "",
                }

                # Line 0: [📍 ]TITLE — COMPANY (LOCATION)
                header = re.sub(r"^📍\s*", "", lines[0]).strip()
                loc_m  = re.search(r"\(([^)]+)\)\s*$", header)
                if loc_m:
                    job["location"] = loc_m.group(1)
                    header = header[:loc_m.start()].strip()
                if " — " in header:
                    t, c = header.split(" — ", 1)
                    job["title"]   = t.strip()
                    job["company"] = c.strip()
                else:
                    job["title"] = header.strip()

                idx = 1
                # Line 1: ↳ REASON
                if idx < len(lines) and lines[idx].startswith("↳"):
                    job["reason"] = lines[idx][1:].strip()
                    idx += 1

                # Lines until Salary: or 🔗 = description snippet
                desc_parts = []
                while idx < len(lines) and not lines[idx].startswith("Salary:") and "🔗" not in lines[idx]:
                    desc_parts.append(lines[idx].strip())
                    idx += 1
                job["description"] = " ".join(desc_parts).strip()

                # Metadata: Salary: X | Posted: Y | Source: Z
                if idx < len(lines) and lines[idx].startswith("Salary:"):
                    for part in lines[idx].split(" | "):
                        part = part.strip()
                        if part.startswith("Salary:"):   job["salary"] = part[7:].strip()
                        elif part.startswith("Posted:"):  job["posted"] = part[7:].strip()
                        elif part.startswith("Source:"):  job["source"] = part[7:].strip()
                    idx += 1

                # URL: 🔗 URL
                if idx < len(lines) and "🔗" in lines[idx]:
                    url = lines[idx].replace("🔗", "").strip()
                    job["url"]    = url
                    job["job_id"] = job_id_from_url(url) if url else ""

                if job["title"] and job["url"]:
                    jobs.append(job)

    return [j for j in jobs if j["job_id"]]

# ── DOC GENERATION ────────────────────────────────────────────────────────────
ALWAYS_EXCL = {"pre-2011-early-career.md", "job-search-strategy.md"}

def _read_doc(filename: str) -> str:
    p = DOCS_DIR / filename
    return p.read_text(encoding="utf-8") if p.exists() else ""


_JD_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                   "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"),
    "Accept": "text/html,application/xhtml+xml,application/json",
    "Accept-Language": "en-US,en;q=0.9",
}


def _html_to_text(s: str) -> str:
    s = re.sub(r"<[^>]+>", " ", s or "")
    s = html.unescape(s)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n\s*\n+", "\n\n", s)
    return s.strip()


def _jsonld_jobposting(page: str) -> str:
    for m in re.finditer(r'<script[^>]*application/ld\+json[^>]*>(.*?)</script>', page, re.S):
        try:
            d = json.loads(m.group(1))
        except Exception:
            continue
        for it in (d if isinstance(d, list) else [d]):
            if isinstance(it, dict) and it.get("@type") == "JobPosting" and it.get("description"):
                return _html_to_text(it["description"])
    return ""


def fetch_job_description(url: str) -> str:
    """Best-effort live fetch of a posting's full description.

    Handles the common ATS JSON APIs (Greenhouse, Ashby, Eightfold) plus a
    generic JSON-LD JobPosting fallback for everything else. Returns plain
    text, or '' if nothing usable was found. Never raises.
    """
    if not url:
        return ""
    try:
        # Greenhouse — direct board URL, or embedded on a company site via gh_jid
        gh_id = re.search(r"gh_jid=(\d+)", url) or \
                (re.search(r"jobs?/(\d+)", url) if "greenhouse" in url else None)
        if gh_id:
            token = None
            mfor = re.search(r"[?&]for=([a-z0-9_]+)", url)
            mtok = re.search(r"greenhouse\.io/(?:embed/job_app\?for=)?([a-z0-9_]+)/jobs?", url, re.I)
            if mfor:
                token = mfor.group(1)
            elif mtok:
                token = mtok.group(1)
            else:
                # Embedded on a company domain: scrape the page for the board token
                try:
                    page = requests.get(url, headers=_JD_HEADERS, timeout=15)
                    t = re.search(r"for=([a-z0-9_]+)", page.text)
                    if t:
                        token = t.group(1)
                except Exception:
                    token = None
            if token:
                r = requests.get(
                    f"https://boards-api.greenhouse.io/v1/boards/{token}/jobs/{gh_id.group(1)}",
                    headers=_JD_HEADERS, timeout=15)
                if r.ok:
                    return _html_to_text(r.json().get("content", ""))

        # Ashby: jobs.ashbyhq.com/{org}/{uuid}
        m = re.search(r"ashbyhq\.com/([^/?]+)/([0-9a-f-]{36})", url, re.I)
        if m:
            r = requests.get(
                f"https://api.ashbyhq.com/posting-api/job-board/{m.group(1)}",
                headers=_JD_HEADERS, timeout=15)
            if r.ok:
                for j in r.json().get("jobs", []):
                    if m.group(2) in json.dumps(j):
                        return _html_to_text(j.get("descriptionPlain") or j.get("descriptionHtml", ""))

        # Eightfold: apply.{company}.com/careers/job/{id}
        m = re.search(r"(apply\.[^/]+)/careers/job/(\d+)", url, re.I)
        if m:
            dom = re.search(r"domain=([^&]+)", url)
            q = f"?domain={dom.group(1)}" if dom else ""
            r = requests.get(f"https://{m.group(1)}/api/apply/v2/jobs/{m.group(2)}{q}",
                             headers=_JD_HEADERS, timeout=15)
            if r.ok:
                return _html_to_text(r.json().get("job_description", ""))

        # Generic: fetch the page and parse a JSON-LD JobPosting
        r = requests.get(url, headers=_JD_HEADERS, timeout=15)
        if r.ok:
            return _jsonld_jobposting(r.text)
    except Exception:
        return ""
    return ""


def ensure_job_description(db, job: dict) -> dict:
    """If a job's stored description is missing or thin, fetch it live and
    persist the result so generation works from the full posting."""
    desc = (job.get("description") or "").strip()
    if len(desc) >= 400:
        return job
    fetched = fetch_job_description(job.get("url") or "")
    if fetched and len(fetched) > len(desc):
        job = dict(job)
        job["description"] = fetched
        db.execute("UPDATE jobs SET description=? WHERE id=?", (fetched, job["id"]))
        db.commit()
    return job


def generate_documents(job: dict, instructions: str = "", fit: Optional[dict] = None) -> dict:
    """Call Claude to generate resume + cover letter markdown.

    If a fit analysis is supplied, its gaps and summary are injected so the
    writing proactively surfaces offsetting experience (without inventing any).
    """
    rules    = _read_doc("resume-generation-rules.md")
    personal = _read_doc("personal-info.md")
    skills   = _read_doc("technical-skills.md")
    edu      = _read_doc("education.md")
    projects = _read_doc("side-projects.md")

    history_parts = [_read_doc(f) for f in JOB_DOCS]
    history = "\n\n---\n\n".join(p for p in history_parts if p)

    regen_block = f"\n\nSPECIAL INSTRUCTIONS:\n{instructions}" if instructions else ""

    fit_block = ""
    if fit:
        summ = (fit.get("summary") or "").strip()
        gaps = (fit.get("gaps") or fit.get("gaps_md") or "").strip()
        if summ or gaps:
            fit_block = f"""

FIT ANALYSIS (already run for this exact candidate-and-job pairing — use it to write smarter):
Overall assessment: {summ}
Known gaps — address each with integrity by surfacing the strongest adjacent or transferable experience that offsets it. NEVER invent experience, and do NOT call the gaps out explicitly in the resume; just make sure the most relevant real experience is front and center:
{gaps}"""

    prompt = f"""You are generating a tailored resume and cover letter for {CANDIDATE_NAME}.

RESUME GENERATION RULES (follow all exactly):
{rules}

ADDITIONAL RULES (these override the rules above where they conflict):
- Output format: clean Markdown — NOT HTML. The output will be converted to DOCX.
- Jobs MUST be listed in strict chronological order, newest to oldest.
- For each job, curate only the most relevant achievements for THIS specific role. Do not dump all bullets — select and tailor.
- LENGTH: target 2 pages, 3 pages absolute maximum. Never produce 4 pages. Be ruthless. Bullets per role: 4-5 for the current/most recent role, 3-4 for mid-career roles, and only 2 for roles older than ~8 years. Keep the summary to 3-4 lines. Selected Projects: at most 2 projects, 1 bullet each. When in doubt, cut the least relevant content rather than keep it.
- BULLET LENGTH: keep every bullet CONCISE — one line, two lines maximum. Lead with the outcome or action, then drop filler clauses ("in order to", "responsible for", long trailing descriptions). A bullet that wraps to three lines is too long; tighten it.
- Resume structure: # Name / contact line / ## Summary / ## Experience / ### Job Title sections / ## Selected Projects (when relevant, see below) / ## Skills / ## Education
- Selected Projects section: AFTER ## Experience, include a ## Selected Projects section IF the target role values hands-on AI, technical depth, data/ML, or a builder profile. Use the same format as jobs: ### Project Name, then a line formatted as **Type / tech** | link-or-status, then 1-2 tailored bullets. Pull ONLY from the SELECTED PROJECTS provided below. Tailor which projects appear to the role. For the Polymarket project, the codebase is private — mention it and its outcomes but never imply the source is public. Omit this whole section for pure process/BA roles where it adds nothing.
- Cover letter: plain paragraphs only — date, greeting, 3-4 body paragraphs, sign-off. No markdown headers.
- Respond ONLY with a valid JSON object — no preamble, no explanation, no markdown fences.
- JSON format: {{"resume": "...", "cover_letter": "..."}}
- Do NOT use **bold** or any other inline emphasis inside bullet points or the summary. Plain text only for all bullets and the summary paragraph. Bold in bullets looks like AI wrote it.
- Use - for bullet points.
- Company name and dates go on the line directly below the ### Job Title line, formatted as: **Company Name** | Location | Start – End
- Skills section: group skills into 4-6 labeled categories, ONE category per line, formatted as: **Category:** skill, skill, skill (for example **API & Integration:** REST, GraphQL, Swagger, Azure APIM). Do not output skills as one long comma list. Tailor the category labels and contents to the target role.

ATS KEYWORD OPTIMIZATION:
Before writing, extract 15-20 key terms from the job description — specific technologies, methodologies, domain concepts, and role-specific phrases the ATS will scan for. Then naturally weave those terms into the resume and cover letter by reformulating existing experience to use the JD's vocabulary. Rules:
- NEVER invent experience. Only rephrase what is already in the work history.
- If the JD says "payment orchestration" and the work history says "payment routing and processing," use "payment orchestration" instead.
- If the JD says "RTP" or "FedNow" and the candidate has real-time payments experience, name those rails explicitly.
- Concentrate keywords in the Summary and the most relevant job bullets.
- Include a Skills section that mirrors the JD's terminology where the candidate has genuine proficiency.
- Do not keyword-stuff — every term must read naturally in context.

GROUNDING — DO NOT FABRICATE (most important rule):
Every bullet and every clause must trace to a specific achievement or metric in the WORK HISTORY below. Build bullets by SELECTING and lightly rephrasing the documented "Key achievements" and "Metrics" — do NOT write net-new accomplishments. If the JD wants something not in the work history, map it to a real documented fact or omit it. BANNED vague filler that signals fabrication: "drove adoption/alignment/cohesion/consensus", "championed", "fostered", "brought structure", "consistent communication", "stakeholder buy-in", "seamless", "synergy". Before returning, re-read EVERY bullet and delete any you cannot trace to a specific fact in the work history below.

ACCURACY GUARDRAILS — ENFORCE BEFORE WRITING:

AI — DAY JOB SCOPE:
The candidate's only documented hands-on AI accomplishment in his day jobs is rolling out GitHub Copilot to one development team, reducing certain stories from 5 story points to 1. Use that fact, scoped to one team, and nothing beyond it. Do not write summary lines implying org-wide or firm-wide AI adoption. Do not frame him as an AI evangelist or AI leader at his employers. Do not say he "builds with AI" at work in any general sense. The scope was one team, one tool.

AI — SIDE PROJECTS ONLY:
Broader AI work (building LLM products, prompt engineering, shipping AI agents) belongs exclusively in the Selected Projects section. Keep it there. Never blend day-job and side-project AI experience into a unified claim or summary line.

TOOL AND SKILL ACCURACY:
Name only tools, platforms, methodologies, and certifications that appear in the candidate's documented technical skills or work history. If the JD lists a tool the candidate does not have, omit it — do not add it. If he has a genuine equivalent, name his real one instead. Never list a tool solely because the JD requests it.

TAILORING DISCIPLINE:
Lead with the candidate's true professional identity and real achievements as the backbone. Tailor by selecting which real experiences to foreground and lightly rephrasing toward the JD's vocabulary. Do not invent a new persona per role. Do not bury his core identity to chase JD keywords. Do not contort the whole narrative to mirror the posting. An over-tailored resume that reads as a different person for every job is a failure.
{regen_block}{fit_block}

CANDIDATE PERSONAL INFO:
{personal}

EDUCATION:
{edu}

TECHNICAL SKILLS:
{skills}

SELECTED PROJECTS (self-built side projects — use for a tailored ## Selected Projects section when the role values AI, ML, technical depth, or a builder profile):
{projects}

WORK HISTORY (newest to oldest — use in this order, do not reorder):
{history}

JOB TO TARGET:
Title:    {job.get('title', '')}
Company:  {job.get('company', '')}
Location: {job.get('location', '')}
Description:
{job.get('description', '(no description provided)')}

Generate the resume and cover letter. Return ONLY the JSON object."""

    result = subprocess.run(
        ["claude", "-p", prompt],
        capture_output=True, text=True, timeout=120, env=_CLI_ENV,
    )
    if result.returncode != 0:
        raise RuntimeError(f"claude CLI error: {result.stderr[:300]}")

    raw   = result.stdout.strip()
    raw   = raw.replace("```json", "").replace("```", "").strip()
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if match:
        data = json.loads(match.group())
        return {
            "resume":       data.get("resume", ""),
            "cover_letter": data.get("cover_letter", ""),
        }
    raise ValueError(f"Could not parse Claude response: {raw[:300]}")


def generate_fit_analysis(job: dict) -> dict:
    """Call Claude to produce a structured fit analysis for a job."""
    personal = _read_doc("personal-info.md")
    skills   = _read_doc("technical-skills.md")
    history_parts = [_read_doc(f) for f in JOB_DOCS]
    history = "\n\n---\n\n".join(p for p in history_parts if p)

    prompt = f"""You are analyzing how well {CANDIDATE_NAME}'s experience matches a job posting.

CANDIDATE PROFILE:
{personal}

TECHNICAL SKILLS:
{skills}

WORK HISTORY (includes Key Achievements and Signature Stories in STAR format):
{history}

JOB POSTING:
Title:    {job.get('title', '')}
Company:  {job.get('company', '')}
Location: {job.get('location', '')}
Description:
{job.get('description', '(no description provided)')}

Produce a fit analysis as a JSON object with these fields:

1. "match_score": A number 1-10. Be honest — a 7 means strong match with minor gaps, a 5 means significant gaps.

2. "matches": A markdown list mapping each major JD requirement to specific experience from the work history. Format each as:
   - **JD Requirement** — matching experience from specific role. Be specific with details.
   Only include requirements where there IS a real match.

3. "gaps": A markdown list of JD requirements the candidate does NOT directly match, with a mitigation strategy for each. Format:
   - **Gap** — mitigation: how adjacent experience or transferable skills could address this.
   If there are no gaps, say "No significant gaps identified."

4. "stories": Pick 3-5 Signature Stories (STAR format) from the work history that are most relevant to THIS job's requirements. For each, write:
   - **Story title** (from which role) — why it's relevant to this JD
   Then reproduce the STAR story verbatim from the work history.

5. "summary": 2-3 sentence overall assessment. What's the strongest selling point? What's the biggest risk?

Return ONLY a valid JSON object with keys: match_score, matches, gaps, stories, summary.
No markdown fences, no preamble."""

    result = subprocess.run(
        ["claude", "-p", prompt],
        capture_output=True, text=True, timeout=120, env=_CLI_ENV,
    )
    if result.returncode != 0:
        raise RuntimeError(f"claude CLI error: {result.stderr[:300]}")

    raw   = result.stdout.strip()
    raw   = raw.replace("```json", "").replace("```", "").strip()
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if match:
        data = json.loads(match.group())
        return {
            "match_score": data.get("match_score", 0),
            "matches":     data.get("matches", ""),
            "gaps":        data.get("gaps", ""),
            "stories":     data.get("stories", ""),
            "summary":     data.get("summary", ""),
        }
    raise ValueError(f"Could not parse fit analysis response: {raw[:300]}")


# ── DOCX EXPORT ───────────────────────────────────────────────────────────────

def _add_formatted_run(paragraph, text: str, size: int = 11,
                        bold: bool = False, italic: bool = False,
                        color: Optional[tuple] = None):
    run = paragraph.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def _inline_runs(paragraph, text: str, size: int = 11, base_color=None):
    """Parse **bold** and *italic* inline markers and add runs."""
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            _add_formatted_run(paragraph, text[last:m.start()], size=size, color=base_color)
        content  = m.group(1) or m.group(2)
        is_bold  = m.group(1) is not None
        _add_formatted_run(paragraph, content, size=size, bold=is_bold,
                           italic=not is_bold, color=base_color)
        last = m.end()
    if last < len(text):
        _add_formatted_run(paragraph, text[last:], size=size, color=base_color)

def _add_rule(doc):
    """Thin horizontal rule above section headers."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "AAAAAA")
    pBdr.append(top)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

BLUE = (31, 73, 125)

def markdown_to_docx(md_text: str) -> Document:
    """Convert resume/cover-letter markdown to a python-docx Document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    # Remove default empty paragraph
    for el in list(doc.element.body):
        doc.element.body.remove(el)

    lines = md_text.split("\n")
    i     = 0
    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()

        if not s:
            i += 1
            continue

        if s.startswith("# "):
            # Name — large, centered, blue
            p   = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formatted_run(p, s[2:].strip(), size=22, bold=True, color=BLUE)
            p.paragraph_format.space_after = Pt(2)

        elif s.startswith("## "):
            # Section header — small caps style, blue, with rule above
            _add_rule(doc)
            p = doc.add_paragraph()
            _add_formatted_run(p, s[3:].strip().upper(), size=10, bold=True, color=BLUE)
            p.paragraph_format.space_after = Pt(3)

        elif s.startswith("### "):
            # Job title
            p = doc.add_paragraph()
            _add_formatted_run(p, s[4:].strip(), size=11, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)

        elif s == "---":
            pass  # skip markdown HRs (we use _add_rule before ## headers)

        elif s.startswith("- "):
            # Bullet
            p = doc.add_paragraph(style="List Bullet")
            _inline_runs(p, s[2:].strip())
            p.paragraph_format.space_after = Pt(2)

        else:
            # Normal paragraph — centered for first 3 lines (contact info), left-aligned after
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 4 else WD_ALIGN_PARAGRAPH.LEFT
            _inline_runs(p, s)
            p.paragraph_format.space_after = Pt(3)

        i += 1

    return doc

# ── STYLED HTML EXPORT (matches davidmyers.work/resume design) ────────────────

RESUME_CSS = """
:root {
  --paper: #f7f4ee; --ink: #1a1714; --ink-soft: #45403a; --ink-faint: #837b6f;
  --line: #ddd5c8; --accent: #0b5d52; --white: #fffdf9;
  --display: "Fraunces", Georgia, serif;
  --body: "Hanken Grotesk", -apple-system, sans-serif;
  --mono: "JetBrains Mono", ui-monospace, monospace;
}
* { box-sizing: border-box; }
body {
  margin: 0; background: var(--paper); color: var(--ink);
  font-family: var(--body); line-height: 1.5; font-size: 15px;
  -webkit-font-smoothing: antialiased;
}
.toolbar {
  position: sticky; top: 0; z-index: 10;
  background: color-mix(in srgb, var(--paper) 90%, transparent);
  backdrop-filter: blur(8px); border-bottom: 1px solid var(--line);
  display: flex; justify-content: space-between; align-items: center;
  padding: 0.85rem clamp(1rem, 4vw, 2rem);
}
.toolbar button {
  font-family: var(--body); font-size: 0.875rem; font-weight: 600;
  border-radius: 100px; padding: 0.55rem 1.2rem; cursor: pointer;
  border: 1px solid var(--accent); background: var(--accent); color: var(--white);
  transition: all .2s ease;
}
.toolbar button:hover { background: #084840; }
.toolbar-hint { font-family: var(--mono); font-size: 0.72rem; color: var(--ink-faint); }
.sheet {
  max-width: 820px; margin: 2rem auto; background: var(--white);
  border: 1px solid var(--line); border-radius: 6px;
  padding: clamp(2rem, 5vw, 3.4rem);
  box-shadow: 0 20px 50px -30px rgba(26,23,20,0.35);
}
.r-name {
  font-family: var(--display); font-weight: 500;
  font-size: clamp(2rem, 5vw, 2.7rem); letter-spacing: -0.02em; margin: 0 0 0.3rem;
}
.r-contact {
  font-family: var(--mono); font-size: 0.78rem; color: var(--ink-soft);
  display: flex; flex-wrap: wrap; gap: 0.4rem 1.1rem;
  padding-bottom: 1.4rem; border-bottom: 2px solid var(--ink); margin-top: 0.6rem;
}
.r-contact a { color: var(--ink-soft); text-decoration: none; }
.r-contact a:hover { color: var(--accent); }
.r-section { margin-top: 1.8rem; }
.r-section > h2 {
  font-family: var(--mono); font-size: 0.76rem; font-weight: 500;
  letter-spacing: 0.08em; text-transform: uppercase; color: var(--accent);
  margin: 0 0 1rem; padding-bottom: 0.45rem; border-bottom: 1px solid var(--line);
}
.r-summary { color: var(--ink-soft); margin: 0; font-size: 1rem; }
.r-job { margin-bottom: 1.4rem; page-break-inside: avoid; }
.r-job:last-child { margin-bottom: 0; }
.r-job-head { display: flex; justify-content: space-between; align-items: baseline; gap: 1rem; flex-wrap: wrap; }
.r-job-title { font-family: var(--display); font-weight: 600; font-size: 1.12rem; margin: 0; letter-spacing: -0.01em; }
.r-job-dates { font-family: var(--mono); font-size: 0.76rem; color: var(--ink-faint); white-space: nowrap; }
.r-job-co { font-weight: 600; color: var(--ink-soft); margin: 0.1rem 0 0.6rem; font-size: 0.92rem; }
.r-job ul, .r-generic-list { margin: 0; padding: 0; list-style: none; display: grid; gap: 0.4rem; }
.r-job li, .r-generic-list li { position: relative; padding-left: 1.1rem; color: var(--ink-soft); font-size: 0.93rem; line-height: 1.5; }
.r-job li::before, .r-generic-list li::before { content: "\\2013"; position: absolute; left: 0; color: var(--accent); }
.r-skills { display: grid; gap: 0.6rem; }
.r-skill-row { display: grid; grid-template-columns: 175px 1fr; gap: 0.8rem; align-items: baseline; }
.r-skill-cat { font-family: var(--mono); font-size: 0.74rem; color: var(--accent); letter-spacing: 0.01em; }
.r-skill-list { color: var(--ink-soft); font-size: 0.92rem; margin: 0; }
.r-cl-body p { color: var(--ink-soft); margin: 0 0 0.8rem; }
@media print {
  @page { margin: 13mm; }
  body { background: #fff; font-size: 10pt; line-height: 1.35; }
  .toolbar { display: none; }
  .sheet { margin: 0; max-width: none; border: none; border-radius: 0; box-shadow: none; padding: 0; background: #fff; }
  .r-name { font-size: 20pt; }
  .r-contact { border-bottom-color: #000; }
  .r-section { margin-top: 11pt; }
  .r-job { margin-bottom: 8pt; }
  .r-job ul, .r-generic-list { gap: 0.3rem; }
  a { color: #000 !important; text-decoration: none; }
  .r-skill-row { grid-template-columns: 150px 1fr; }
}
@media (max-width: 560px) { .r-skill-row { grid-template-columns: 1fr; gap: 0.15rem; } }
"""

_HTML_SHELL = """<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>__TITLE__</title>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,500;9..144,600&family=Hanken+Grotesk:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet"/>
<style>__CSS__</style></head>
<body>
<div class="toolbar">
  <span class="toolbar-hint">Tip: print &rarr; "Save as PDF" (margins: Default, background graphics: on)</span>
  <button class="print" type="button" onclick="window.print()">Save as PDF</button>
</div>
<article class="sheet">__BODY__</article>
</body></html>"""


def _inline_html(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    return text


def _strip_md(s: str) -> str:
    return s.replace("**", "").replace("*", "").strip()


def _render_contact(lines: list[str]) -> str:
    parts = []
    for ln in lines:
        for piece in ln.split("|"):
            piece = _strip_md(piece)
            if piece:
                parts.append(piece)
    spans = []
    for p in parts:
        safe = _inline_html(p)
        low  = p.lower()
        if "@" in p and " " not in p:
            spans.append(f'<a href="mailto:{p}">{safe}</a>')
        elif low.startswith("http") or "linkedin.com" in low or low.startswith("www.") or ".work" in low or ".io" in low:
            href = p if low.startswith("http") else "https://" + p
            spans.append(f'<a href="{href}" target="_blank" rel="noopener">{safe}</a>')
        else:
            spans.append(f"<span>{safe}</span>")
    return "\n".join(spans)


def _parse_job_meta(meta: str):
    """`**Company** | Location | Start - End` -> (company, location, dates)."""
    parts = [_strip_md(p) for p in meta.split("|")]
    parts = [p for p in parts if p]
    if len(parts) >= 3:
        return parts[0], " · ".join(parts[1:-1]), parts[-1]
    if len(parts) == 2:
        if any(ch.isdigit() for ch in parts[1]):
            return parts[0], "", parts[1]
        return parts[0], parts[1], ""
    if parts:
        return parts[0], "", ""
    return "", "", ""


def _render_jobs(body: list[str]) -> str:
    jobs, cur = [], None
    for raw in body:
        s = raw.strip()
        if s.startswith("### "):
            if cur:
                jobs.append(cur)
            cur = {"title": s[4:].strip(), "meta": "", "bullets": []}
        elif cur is not None:
            if s.startswith("- "):
                cur["bullets"].append(s[2:].strip())
            elif s and not cur["meta"] and not s.startswith("#"):
                cur["meta"] = s
    if cur:
        jobs.append(cur)

    out = []
    for job in jobs:
        company, location, dates = _parse_job_meta(job["meta"])
        co = f"{company} · {location}" if company and location else (company or location)
        out.append('<div class="r-job"><div class="r-job-head">')
        out.append(f'<h3 class="r-job-title">{_inline_html(job["title"])}</h3>')
        if dates:
            out.append(f'<span class="r-job-dates">{_inline_html(dates)}</span>')
        out.append("</div>")
        if co:
            out.append(f'<p class="r-job-co">{_inline_html(co)}</p>')
        if job["bullets"]:
            out.append("<ul>")
            out += [f"<li>{_inline_html(b)}</li>" for b in job["bullets"]]
            out.append("</ul>")
        out.append("</div>")
    return "\n".join(out)


def _split_skill(item: str):
    for pat in (r"\*\*(.+?)\*\*\s*[:—–-]\s*(.+)", r"\*\*(.+?)\*\*\s+(.+)",
                r"([A-Za-z][A-Za-z &/]{2,30}?):\s+(.+)"):
        m = re.match(pat, item)
        if m:
            return m.group(1).strip(), m.group(2).strip()
    return "", item


def _render_skills(body: list[str]) -> str:
    out = ['<div class="r-skills">']
    for raw in body:
        s = raw.strip()
        if not s:
            continue
        item = s[2:].strip() if s.startswith("- ") else s
        cat, lst = _split_skill(item)
        if cat:
            out.append('<div class="r-skill-row">')
            out.append(f'<span class="r-skill-cat">{_inline_html(cat)}</span>')
            out.append(f'<p class="r-skill-list">{_inline_html(lst)}</p>')
            out.append("</div>")
        else:
            out.append(f'<p class="r-skill-list">{_inline_html(item)}</p>')
    out.append("</div>")
    return "\n".join(out)


def _render_generic(body: list[str], summary: bool = False) -> str:
    out, bullets, para = [], [], []

    def flush_para():
        if para:
            cls = ' class="r-summary"' if summary else ""
            out.append(f'<p{cls}>{_inline_html(" ".join(para).strip())}</p>')
            para.clear()

    def flush_bullets():
        if bullets:
            out.append('<ul class="r-generic-list">')
            out.extend(f"<li>{_inline_html(b)}</li>" for b in bullets)
            out.append("</ul>")
            bullets.clear()

    for raw in body:
        s = raw.strip()
        if not s:
            flush_para()
        elif s.startswith("- "):
            flush_para()
            bullets.append(s[2:].strip())
        else:
            flush_bullets()
            para.append(s)
    flush_para()
    flush_bullets()
    return "\n".join(out)


def _md_resume_body(md_text: str) -> str:
    lines = md_text.split("\n")
    name, contact, i = "", [], 0
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("# ") and not s.startswith("## "):
            name = s[2:].strip()
            i += 1
            break
        i += 1
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("## ") or s.startswith("### "):
            break
        if s and not s.startswith("#"):
            contact.append(s)
        i += 1

    sections, cur_title, cur_lines = [], None, []
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("## ") and not s.startswith("### "):
            if cur_title is not None:
                sections.append((cur_title, cur_lines))
            cur_title, cur_lines = s[3:].strip(), []
        elif cur_title is not None:
            cur_lines.append(lines[i])
        i += 1
    if cur_title is not None:
        sections.append((cur_title, cur_lines))

    out = ["<header>", f'<h1 class="r-name">{_inline_html(name)}</h1>']
    if contact:
        out.append(f'<div class="r-contact">{_render_contact(contact)}</div>')
    out.append("</header>")

    for title, sbody in sections:
        out.append('<section class="r-section">')
        out.append(f"<h2>{_inline_html(title)}</h2>")
        if any(l.strip().startswith("### ") for l in sbody):
            out.append(_render_jobs(sbody))
        elif title.strip().lower().startswith("skill"):
            out.append(_render_skills(sbody))
        else:
            out.append(_render_generic(sbody, summary=(title.strip().lower() == "summary")))
        out.append("</section>")
    return "\n".join(out)


def _md_coverletter_body(md_text: str) -> str:
    lines = md_text.split("\n")

    # Optional letterhead: a leading `# Name` heading plus an immediately
    # following contact line (containing | or @) render as a styled header,
    # matching the resume. Everything after is normal letter body.
    i = 0
    name, contact = "", []
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("# ") and not s.startswith("## "):
            name = s[2:].strip()
            i += 1
            break
        if s and not s.startswith("#"):
            break  # body starts before any name heading; no letterhead
        i += 1
    if name:
        j = i
        while j < len(lines) and not lines[j].strip():
            j += 1
        if j < len(lines):
            cand = lines[j].strip()
            if ("|" in cand or "@" in cand) and not cand.lower().startswith("dear"):
                contact.append(cand)
                i = j + 1

    out = []
    if name:
        out.append("<header>")
        out.append(f'<h1 class="r-name">{_inline_html(name)}</h1>')
        if contact:
            out.append(f'<div class="r-contact">{_render_contact(contact)}</div>')
        out.append("</header>")

    out.append('<div class="r-cl-body">')
    para = []

    def flush():
        # Blank line = new paragraph; single newline within a block = line break.
        # Keeps body paragraphs reflowing while preserving sign-off/address blocks.
        if para:
            out.append("<p>" + "<br>".join(_inline_html(p) for p in para) + "</p>")
            para.clear()

    for raw in lines[i:]:
        s = raw.strip()
        if not s:
            flush()
        elif s.startswith("#"):
            continue
        else:
            para.append(s)
    flush()
    out.append("</div>")
    return "\n".join(out)


def markdown_to_html(md_text: str, kind: str = "resume", title: str = "Resume") -> str:
    body = _md_coverletter_body(md_text) if kind == "coverletter" else _md_resume_body(md_text)
    return (_HTML_SHELL
            .replace("__TITLE__", title)
            .replace("__CSS__", RESUME_CSS)
            .replace("__BODY__", body))


def html_to_pdf_bytes(html: str) -> bytes:
    """Render styled HTML to a PDF with headless Chromium (Playwright).
    Matches the browser 'Save as PDF' output exactly: same engine, fonts, and
    @media print rules. The toolbar is hidden by the print stylesheet."""
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(args=["--no-sandbox"])
        try:
            page = browser.new_page()
            page.set_content(html, wait_until="load")
            try:
                page.evaluate("() => document.fonts.ready")  # let web fonts settle
            except Exception:
                pass
            return page.pdf(
                print_background=True,
                prefer_css_page_size=True,  # honor the @page { margin } in print CSS
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
            )
        finally:
            browser.close()


# ── STORY BANK PARSER ─────────────────────────────────────────────────────────

def _parse_stories_from_docs() -> list[dict]:
    """Extract Signature Story (STAR) sections from all job docs."""
    stories = []
    for filename in JOB_DOCS:
        text = _read_doc(filename)
        if not text:
            continue

        # Extract company/title from frontmatter
        company = title = ""
        for line in text.split("\n"):
            if line.startswith("company:"):
                company = line.split(":", 1)[1].strip().strip('"')
            elif line.startswith("title:"):
                title = line.split(":", 1)[1].strip().strip('"')

        # Find Signature story section
        sig_match = re.search(
            r"## Signature stor(?:y|ies)\s*\(STAR(?:\s+format)?\)\s*\n(.*?)(?=\n## |\Z)",
            text, re.DOTALL,
        )
        if not sig_match:
            continue

        body = sig_match.group(1).strip()

        # Check for ### sub-stories (multiple stories per role)
        sub_stories = re.split(r"### (.+)\n", body)
        if len(sub_stories) > 1:
            # sub_stories = ['', 'Story 1 — title', 'content', 'Story 2 — title', 'content', ...]
            for i in range(1, len(sub_stories), 2):
                story_title = sub_stories[i].strip()
                story_body  = sub_stories[i + 1].strip() if i + 1 < len(sub_stories) else ""
                stories.append({
                    "title":    story_title,
                    "role":     title,
                    "company":  company,
                    "filename": filename,
                    "body":     story_body,
                })
        else:
            # Single story
            stories.append({
                "title":    f"{title} — {company}",
                "role":     title,
                "company":  company,
                "filename": filename,
                "body":     body,
            })

    return stories

# ── ROUTES ────────────────────────────────────────────────────────────────────

@app.route("/")
def board():
    db = get_db()
    rows = data.get_board_jobs(db)
    stale_ids = data.get_stale_job_ids(db)

    jobs_by_status = {s: [] for s in STATUSES}
    for row in rows:
        status = row["status"] if row["status"] in STATUSES else "New"
        job = dict(row)
        job["is_stale"] = row["id"] in stale_ids
        jobs_by_status[status].append(job)

    return render_template(
        "board.html",
        jobs_by_status=jobs_by_status,
        app_statuses=APP_STATUSES,
        end_statuses=END_STATUSES,
        statuses=STATUSES,
        status_colors=STATUS_COLORS,
    )


@app.route("/interviews")
def interviews():
    db   = get_db()
    rows = data.get_interview_jobs(db)

    jobs = []
    for row in rows:
        job = dict(row)
        job["recruiter_name"] = data.get_recruiter_name(db, job["id"])
        jobs.append(job)

    return render_template(
        "interviews.html",
        jobs=jobs,
        status_colors=STATUS_COLORS,
    )


@app.route("/radar")
@app.route("/radar/<date>/<slot>")
def radar(date=None, slot=None):
    reports = list_reports()
    if not reports:
        return render_template("radar.html", jobs=[], reports=reports, current=None)

    if date and slot:
        path     = OUTPUT_DIR / f"{date}-{slot.lower()}.md"
        filename = path.name
    else:
        path     = Path(reports[0]["path"])
        filename = reports[0]["filename"]
        date     = reports[0]["date"]
        slot     = reports[0]["slot"]

    jobs = parse_report(str(path)) if path.exists() else []

    show_dismissed = request.args.get("dismissed") == "1"

    db             = get_db()
    saved_ids      = data.get_all_job_ids(db)
    dismissed_ids  = data.get_dismissed_ids(db)
    reviewed_files = data.get_reviewed_filenames(db)
    comments       = data.get_radar_comments(db)

    company_signals = data.get_company_signals(db)

    for job in jobs:
        job["is_saved"]     = job["job_id"] in saved_ids
        job["is_dismissed"] = job["job_id"] in dismissed_ids
        job["comment"]      = comments.get(job["job_id"], "")
        key = (job.get("company") or "").lower().strip()
        job["company_signal"] = company_signals.get(key)

    if not show_dismissed:
        jobs = [j for j in jobs if not j["is_dismissed"]]

    is_reviewed = filename in reviewed_files
    for r in reports:
        r["is_reviewed"] = r["filename"] in reviewed_files

    # Load filtered jobs for this report (if a matching run exists)
    filtered_grouped = {}
    run = data.get_run_by_report_file(db, filename)
    if run:
        for fj in data.get_filtered_jobs(db, run["id"]):
            filtered_grouped.setdefault(fj["filter_name"], []).append(dict(fj))

    current = {"filename": filename, "date": date, "slot": slot}
    return render_template("radar.html", jobs=jobs, reports=reports, current=current,
                           show_dismissed=show_dismissed, dismissed_count=len(dismissed_ids),
                           is_reviewed=is_reviewed, filtered_grouped=filtered_grouped)


@app.route("/stories")
def story_bank():
    stories = _parse_stories_from_docs()

    # Collect fit analysis stories to show which jobs each story was surfaced for
    db = get_db()
    analyses = data.get_fit_analyses_with_stories(db)

    # For each doc story, find which fit analyses referenced it (by keyword match)
    for story in stories:
        story["used_in"] = []
        # Match by a key phrase from the story body (first 60 chars of the Situation line)
        situation_line = ""
        for line in story["body"].split("\n"):
            if line.startswith("**Situation:**"):
                situation_line = line.replace("**Situation:**", "").strip()[:60]
                break
        if situation_line:
            for a in analyses:
                if a["stories_md"] and situation_line[:40] in a["stories_md"]:
                    story["used_in"].append({
                        "title":   a["title"],
                        "company": a["company"],
                    })

    return render_template("stories.html", stories=stories)


@app.route("/health")
def health():
    db = get_db()
    latest = data.get_latest_run(db)
    if not latest:
        return render_template("health.html", run=None, stats=[], filter_stats=[],
                               filter_jobs={}, skip_reasons={}, history=[])

    run = dict(latest)
    stats = [dict(r) for r in data.get_source_stats(db, run["id"])]

    prev_stats = data.get_previous_run_stats(db, run["id"])
    prev_map = {r["source"]: r["raw_count"] for r in prev_stats}

    for s in stats:
        prev = prev_map.get(s["source"])
        if prev is None:
            s["trend"] = "new"
        elif s["raw_count"] > prev:
            s["trend"] = "up"
        elif s["raw_count"] < prev:
            s["trend"] = "down"
        else:
            s["trend"] = "flat"

    filter_stats = [dict(r) for r in data.get_filter_stats(db, run["id"])]
    history = [dict(r) for r in data.get_recent_runs(db, days=14)]

    filter_jobs: dict[str, list] = {}
    for row in data.get_filtered_jobs(db, run["id"]):
        filter_jobs.setdefault(row["filter_name"], []).append(dict(row))

    from filters import SKIP_REASONS  # local import — filters loads config

    return render_template("health.html", run=run, stats=stats,
                           filter_stats=filter_stats, history=history,
                           filter_jobs=filter_jobs, skip_reasons=SKIP_REASONS)


@app.route("/skipped")
def skipped():
    db = get_db()
    latest = data.get_latest_run(db)
    if not latest:
        return render_template("skipped.html", run=None, groups=[], total=0)

    run = dict(latest)
    rows = data.get_filtered_jobs(db, run["id"])

    from filters import SKIP_REASONS, explain_skip

    grouped: dict[str, list[dict]] = {}
    for row in rows:
        grouped.setdefault(row["filter_name"], []).append(dict(row))

    groups = [
        {
            "filter_name": name,
            "reason": explain_skip(name),
            "jobs": jobs,
        }
        for name, jobs in sorted(grouped.items(), key=lambda kv: -len(kv[1]))
    ]

    return render_template("skipped.html", run=run, groups=groups,
                           total=len(rows), skip_reasons=SKIP_REASONS)


@app.route("/portals")
def portals():
    db = get_db()
    scan = data.get_latest_scan(db)

    jobs = []
    if scan:
        scan = dict(scan)
        rows = data.get_scan_results(db, scan["id"])

        saved_ids     = data.get_all_job_ids(db)
        dismissed_ids = data.get_dismissed_ids(db)

        remote_signals = ("remote", "work from home", "wfh", "distributed", "virtual")
        for row in rows:
            j = dict(row)
            j["is_saved"]     = j["job_id"] in saved_ids
            j["is_dismissed"] = j["job_id"] in dismissed_ids
            loc = (j.get("location") or "").lower()
            if any(t in loc for t in HOME_METRO_TERMS):
                j["loc_tag"] = "local"
            elif any(s in loc for s in remote_signals):
                j["loc_tag"] = "remote"
            else:
                j["loc_tag"] = "other"
            jobs.append(j)

    loc_counts = {"local": 0, "remote": 0, "other": 0}
    for j in jobs:
        loc_counts[j.get("loc_tag", "other")] += 1

    return render_template(
        "portals.html",
        scan=scan,
        jobs=jobs,
        company_count=len(set(j["company"] for j in jobs)),
        loc_counts=loc_counts,
        local_label=f"{HOME_CITY} area",
    )


@app.route("/portals/scan", methods=["POST"])
def run_portal_scan():
    try:
        results = portal_scan_all()
        db = get_db()
        companies_hit = len(set(j["company"] for j in results))
        scan_id = data.insert_scan(db, len(results), companies_hit)
        data.insert_scan_results(db, scan_id, results)
        return jsonify({"ok": True, "found": len(results), "companies": companies_hit})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/jobs/save", methods=["POST"])
def save_job():
    payload = request.json or {}
    job_id  = payload.get("job_id")
    if not job_id:
        return jsonify({"ok": False, "error": "missing job_id"}), 400
    db = get_db()
    try:
        data.save_job(db, job_id, payload)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/radar/reviewed", methods=["POST"])
def mark_radar_reviewed():
    filename = (request.json or {}).get("filename", "")
    if not filename:
        return jsonify({"ok": False, "error": "missing filename"}), 400
    db = get_db()
    data.mark_reviewed(db, filename)
    return jsonify({"ok": True})


@app.route("/radar/dismiss/<job_id>", methods=["POST"])
def dismiss_radar_job(job_id):
    payload = request.json or {}
    undo    = payload.get("undo", False)
    company = payload.get("company", "")
    db = get_db()
    if undo:
        data.undismiss_job(db, job_id)
    else:
        data.dismiss_job(db, job_id)
        if company:
            data.increment_company_signal(db, company, "dismiss")
    return jsonify({"ok": True})


@app.route("/radar/comments/<job_id>", methods=["POST"])
def save_radar_comment(job_id):
    body = (request.json or {}).get("body", "").strip()
    db   = get_db()
    data.save_radar_comment(db, job_id, body)
    return jsonify({"ok": True})


@app.route("/company/boost", methods=["POST"])
def boost_company():
    company = (request.json or {}).get("company", "").strip()
    if not company:
        return jsonify({"ok": False, "error": "missing company"}), 400
    db = get_db()
    data.set_company_boost(db, company)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/move", methods=["POST"])
def move_job(job_id):
    payload = request.json or {}
    status  = payload.get("status") or request.form.get("status", "")
    if status not in STATUSES:
        return jsonify({"ok": False, "error": "invalid status"}), 400
    db = get_db()
    if status == "Ready":
        check = data.check_ready_requirements(db, job_id)
        if not check["ok"]:
            return jsonify({"ok": False, "error": "Missing requirements", "missing": check["missing"]}), 400
    if status in ("Rejected", "Passed"):
        job = data.get_job(db, job_id)
        if job and job["company"]:
            data.increment_company_signal(db, job["company"], "reject")
    data.move_job(db, job_id, status)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/notes", methods=["POST"])
def add_note(job_id):
    payload = request.json or {}
    body = (payload.get("body") or request.form.get("body", "")).strip()
    if not body:
        return jsonify({"ok": False, "error": "empty note"}), 400
    db = get_db()
    data.add_note(db, job_id, body)
    notes = data.get_notes(db, job_id)
    return jsonify({"ok": True, "notes": [dict(n) for n in notes]})


@app.route("/jobs/<job_id>")
def job_detail(job_id):
    db  = get_db()
    job = data.get_job(db, job_id)
    if not job:
        return "Job not found", 404
    notes     = data.get_notes(db, job_id)
    docs      = data.get_documents(db, job_id)
    apply_url = data.get_apply_url(db, job_id)
    fit       = data.get_fit_analysis(db, job_id)

    contacts = offer = rounds = None
    prep_docs = prep_files = None
    if dict(job)["status"] in ("Interviewing", "Offer"):
        contacts = data.get_contacts(db, job_id)
        rounds   = data.get_rounds(db, job_id)
        prep_docs = data.get_prep_docs(db, job_id)
        linked    = {p["filename"] for p in prep_docs}
        prep_files = [f for f in _list_prep_files() if f not in linked]
    if dict(job)["status"] == "Offer":
        offer = data.get_offer(db, job_id)

    return render_template(
        "job_detail.html",
        job=dict(job),
        notes=[dict(n) for n in notes],
        docs=[dict(d) for d in docs],
        apply_url=apply_url,
        fit=fit,
        contacts=contacts,
        rounds=rounds,
        prep_docs=prep_docs,
        prep_files=prep_files,
        offer=offer,
        statuses=STATUSES,
        status_colors=STATUS_COLORS,
    )


@app.route("/jobs/<job_id>/apply_url", methods=["POST"])
def save_apply_url(job_id):
    url = (request.json or {}).get("url", "").strip()
    if url and not url.lower().startswith(("http://", "https://", "mailto:")):
        return jsonify({"ok": False, "error": "URL must start with http://, https://, or mailto:"}), 400
    db = get_db()
    if url:
        data.save_apply_url(db, job_id, url)
    else:
        data.delete_apply_url(db, job_id)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/contacts", methods=["POST"])
def add_contact(job_id):
    payload = request.json or {}
    name = payload.get("name", "").strip()
    if not name:
        return jsonify({"ok": False, "error": "name required"}), 400
    db = get_db()
    new_id = data.add_contact(db, job_id, payload)
    return jsonify({"ok": True, "id": new_id})


@app.route("/jobs/<job_id>/contacts/<int:contact_id>", methods=["DELETE"])
def delete_contact(job_id, contact_id):
    db = get_db()
    data.delete_contact(db, contact_id, job_id)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/rounds", methods=["POST"])
def add_round(job_id):
    payload = request.json or {}
    db = get_db()
    data.add_round(db, job_id, payload)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/rounds/<int:round_id>", methods=["POST"])
def update_round(job_id, round_id):
    payload = request.json or {}
    db = get_db()
    data.update_round(db, round_id, job_id, payload)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/rounds/<int:round_id>", methods=["DELETE"])
def delete_round(job_id, round_id):
    db = get_db()
    data.delete_round(db, round_id, job_id)
    return jsonify({"ok": True})


# ── Interview Prep Docs ───────────────────────────────────────────────────────

def _list_prep_files():
    """Markdown prep files available to link, newest first."""
    if not INTERVIEW_PREP_DIR.exists():
        return []
    files = sorted(INTERVIEW_PREP_DIR.glob("*.md"),
                   key=lambda p: p.stat().st_mtime, reverse=True)
    return [p.name for p in files]


def _prep_title_from_filename(fname):
    return fname.rsplit(".", 1)[0].replace("-", " ").replace("_", " ").strip()


def _render_prep_html(md_text, title):
    import markdown as _md
    body = _md.markdown(md_text, extensions=["tables", "fenced_code", "sane_lists"])
    safe_title = (title or "Interview Prep").replace("<", "&lt;").replace(">", "&gt;")
    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>{safe_title}</title>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,500;9..144,600&family=Hanken+Grotesk:wght@400;500;600;700&display=swap" rel="stylesheet"/>
<style>
  body{{max-width:760px;margin:0 auto;padding:1.5rem 1.5rem 4rem;font-family:'Hanken Grotesk',-apple-system,sans-serif;line-height:1.55;color:#1a1714;background:#f7f4ee;}}
  h1,h2,h3{{font-family:'Fraunces',Georgia,serif;line-height:1.2;letter-spacing:-0.01em;}}
  h1{{font-size:1.8rem;margin:0 0 1rem;}}
  h2{{font-size:1.25rem;margin:1.9rem 0 .6rem;border-bottom:1px solid #ddd5c8;padding-bottom:.3rem;color:#0b5d52;}}
  h3{{font-size:1.05rem;margin:1.2rem 0 .4rem;}}
  a{{color:#0b5d52;}} code{{background:#efe9df;padding:.1em .35em;border-radius:4px;font-size:.9em;}}
  ul,ol{{padding-left:1.3rem;}} li{{margin:.25rem 0;}}
  blockquote{{border-left:3px solid #0b5d52;margin:.8rem 0;padding:.2rem 0 .2rem 1rem;color:#45403a;}}
  .bar{{display:flex;justify-content:space-between;align-items:center;margin-bottom:1.4rem;font-size:.85rem;}}
  .bar button{{font:inherit;font-weight:600;border:1px solid #0b5d52;background:#0b5d52;color:#fff;border-radius:100px;padding:.5rem 1.1rem;cursor:pointer;}}
  .bar a{{text-decoration:none;color:#45403a;}}
  @media print{{.bar{{display:none;}}body{{background:#fff;}}}}
</style></head>
<body>
<div class="bar"><a href="javascript:history.back()">&larr; Back</a><button onclick="window.print()">Save as PDF</button></div>
{body}
</body></html>"""


@app.route("/jobs/<job_id>/prep", methods=["POST"])
def add_prep(job_id):
    payload  = request.json or {}
    filename = (payload.get("filename") or "").strip()
    if filename not in _list_prep_files():
        return jsonify({"ok": False, "error": "unknown prep file"}), 400
    title = (payload.get("title") or _prep_title_from_filename(filename)).strip()
    db = get_db()
    new_id = data.add_prep_doc(db, job_id, filename, title)
    return jsonify({"ok": True, "id": new_id, "title": title, "filename": filename})


@app.route("/jobs/<job_id>/prep/<int:prep_id>", methods=["DELETE"])
def delete_prep(job_id, prep_id):
    db = get_db()
    data.delete_prep_doc(db, prep_id, job_id)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/prep/<int:prep_id>/view")
def view_prep(job_id, prep_id):
    db  = get_db()
    doc = data.get_prep_doc(db, prep_id, job_id)
    if not doc:
        return "Not found", 404
    fname = doc["filename"]
    if fname not in _list_prep_files():
        return "Prep file is no longer on disk", 404
    text = (INTERVIEW_PREP_DIR / fname).read_text(encoding="utf-8", errors="replace")
    return _render_prep_html(text, doc.get("title") or fname)


@app.route("/jobs/<job_id>/offer", methods=["POST"])
def save_offer(job_id):
    payload = request.json or {}
    db = get_db()
    data.save_offer(db, job_id, payload)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/generate", methods=["POST"])
def generate(job_id):
    db  = get_db()
    job = data.get_job(db, job_id)
    if not job:
        return jsonify({"ok": False, "error": "job not found"}), 404

    payload      = request.json or {}
    instructions = payload.get("instructions", "").strip()

    try:
        # 1) Make sure we're working from the full, live posting — not a thin
        #    or stale stored description.
        job = ensure_job_description(db, dict(job))

        # 2) Run (or reuse) a fit analysis before writing, so generation is
        #    gap-aware. Best-effort: never let this block document generation.
        fit = None
        try:
            existing = data.get_fit_analysis(db, job_id)
            if existing:
                fit = dict(existing)
            else:
                fit = generate_fit_analysis(job)
                data.save_fit_analysis(db, job_id, fit)
        except Exception:
            fit = None

        # 3) Generate the documents, informed by the fit analysis.
        result = generate_documents(job, instructions, fit)
        max_v  = data.get_max_doc_version(db, job_id)
        data.insert_document(db, job_id, max_v + 1,
                             result["resume"], result["cover_letter"], instructions)
        db.execute("UPDATE jobs SET status='Drafting' WHERE id=? AND status IN ('New','Reviewing')", (job_id,))
        db.commit()
        doc_id = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]
        return jsonify({
            "ok": True, "doc_id": doc_id, "version": max_v + 1,
            "fit_score": (fit or {}).get("match_score"),
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/jobs/<job_id>/documents/latest")
def latest_document(job_id):
    """Latest document id for a job. Lets the UI recover when a slow generation
    is cut off by a gateway/tunnel timeout but still completes server-side."""
    db  = get_db()
    row = db.execute(
        "SELECT id, version FROM documents WHERE job_id=? ORDER BY version DESC LIMIT 1",
        (job_id,),
    ).fetchone()
    if row:
        return jsonify({"ok": True, "doc_id": row["id"], "version": row["version"]})
    return jsonify({"ok": True, "doc_id": None, "version": 0})


@app.route("/jobs/<job_id>/fit-analysis", methods=["POST"])
def run_fit_analysis(job_id):
    db  = get_db()
    job = data.get_job(db, job_id)
    if not job:
        return jsonify({"ok": False, "error": "job not found"}), 404
    try:
        result = generate_fit_analysis(dict(job))
        data.save_fit_analysis(db, job_id, result)
        return jsonify({"ok": True, "score": result["match_score"]})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/jobs/<job_id>/documents/<int:doc_id>")
def view_draft(job_id, doc_id):
    db  = get_db()
    job = data.get_job(db, job_id)
    doc = data.get_document(db, doc_id, job_id)
    if not job or not doc:
        return "Not found", 404
    all_docs = data.get_all_doc_versions(db, job_id)
    return render_template(
        "draft.html",
        job=dict(job),
        doc=dict(doc),
        all_docs=[dict(d) for d in all_docs],
        statuses=STATUSES,
    )


@app.route("/jobs/<job_id>/documents/<int:doc_id>/save", methods=["POST"])
def save_document(job_id, doc_id):
    payload    = request.json or {}
    resume_md  = payload.get("resume_md", "")
    cl_md      = payload.get("coverletter_md", "")
    mark_final = bool(payload.get("mark_final"))
    db = get_db()
    data.save_document(db, doc_id, job_id, resume_md, cl_md, mark_final)
    return jsonify({"ok": True})


@app.route("/jobs/<job_id>/documents/<int:doc_id>/download/<doc_type>")
def download_doc(job_id, doc_id, doc_type):
    db  = get_db()
    job = data.get_job(db, job_id)
    doc = data.get_document(db, doc_id, job_id)
    if not job or not doc:
        return "Not found", 404

    if doc_type == "resume":
        md_text  = doc["resume_md"] or ""
        prefix   = "Resume"
    else:
        md_text  = doc["coverletter_md"] or ""
        prefix   = "CoverLetter"

    safe_company = re.sub(r"[^\w]", "_", job["company"] or "")[:30]
    safe_title   = re.sub(r"[^\w]", "_", job["title"]   or "")[:30]
    filename     = f"{prefix}_{safe_company}_{safe_title}_v{doc['version']}.docx"
    tmp_path     = EXPORT_DIR / filename

    word_doc = markdown_to_docx(md_text)
    word_doc.save(str(tmp_path))

    return send_file(
        str(tmp_path),
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/jobs/<job_id>/documents/<int:doc_id>/styled/<doc_type>")
def view_styled(job_id, doc_id, doc_type):
    """Serve a styled HTML resume/cover letter (matches davidmyers.work design).
    Opens inline so the user can print to PDF."""
    db  = get_db()
    job = data.get_job(db, job_id)
    doc = data.get_document(db, doc_id, job_id)
    if not job or not doc:
        return "Not found", 404

    if doc_type == "resume":
        return markdown_to_html(doc["resume_md"] or "", kind="resume",
                                title=f"{job['company']} — Resume")
    return markdown_to_html(doc["coverletter_md"] or "", kind="coverletter",
                            title=f"{job['company']} — Cover Letter")


@app.route("/jobs/<job_id>/documents/<int:doc_id>/pdf/<doc_type>")
def download_pdf(job_id, doc_id, doc_type):
    """One-click styled PDF (headless Chromium render of the styled HTML)."""
    db  = get_db()
    job = data.get_job(db, job_id)
    doc = data.get_document(db, doc_id, job_id)
    if not job or not doc:
        return "Not found", 404

    if doc_type == "resume":
        html   = markdown_to_html(doc["resume_md"] or "", kind="resume",
                                  title=f"{job['company']} — Resume")
        prefix = "Resume"
    else:
        html   = markdown_to_html(doc["coverletter_md"] or "", kind="coverletter",
                                  title=f"{job['company']} — Cover Letter")
        prefix = "CoverLetter"

    safe_company = re.sub(r"[^\w]", "_", job["company"] or "")[:30]
    safe_title   = re.sub(r"[^\w]", "_", job["title"]   or "")[:30]
    filename     = f"{prefix}_{safe_company}_{safe_title}_v{doc['version']}.pdf"
    tmp_path     = EXPORT_DIR / filename

    try:
        pdf_bytes = html_to_pdf_bytes(html)
    except Exception as e:
        return jsonify({"ok": False, "error": f"PDF render failed: {e}"}), 500

    with open(tmp_path, "wb") as f:
        f.write(pdf_bytes)

    return send_file(str(tmp_path), as_attachment=True,
                     download_name=filename, mimetype="application/pdf")


@app.route("/jobs/<job_id>/delete", methods=["POST"])
def delete_job(job_id):
    db = get_db()
    data.delete_job(db, job_id)
    return jsonify({"ok": True})


# ── MAIN ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    init_db()
    print("─" * 50)
    print("  Job Search Dashboard")
    print("  http://localhost:5000")
    print("─" * 50)
    app.run(debug=False, host="127.0.0.1", port=5000, threaded=True)
